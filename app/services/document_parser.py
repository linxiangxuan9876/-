"""
文档解析服务 - 支持多种格式的文档解析
"""
import os
import re
from typing import Dict, List, Optional, Tuple
from pathlib import Path


class DocumentParser:
    """文档解析器 - 支持PDF、Word、TXT等格式"""

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt', '.md'}

    @classmethod
    def parse(cls, file_path: str) -> Dict:
        """
        解析文档，返回结构化数据
        Returns:
            {
                "success": bool,
                "content": str,  # 文档全文
                "text_blocks": List[str],  # 文本块列表
                "metadata": Dict,  # 文档元数据
                "error": str  # 错误信息（如果有）
            }
        """
        file_path = Path(file_path)
        ext = file_path.suffix.lower()

        if ext not in cls.SUPPORTED_EXTENSIONS:
            return {
                "success": False,
                "content": "",
                "text_blocks": [],
                "metadata": {},
                "error": f"不支持的文件格式: {ext}"
            }

        try:
            if ext == '.pdf':
                return cls._parse_pdf(file_path)
            elif ext in ['.docx', '.doc']:
                return cls._parse_word(file_path)
            elif ext in ['.txt', '.md']:
                return cls._parse_text(file_path)
            else:
                return {
                    "success": False,
                    "content": "",
                    "text_blocks": [],
                    "metadata": {},
                    "error": f"未实现的解析器: {ext}"
                }
        except Exception as e:
            return {
                "success": False,
                "content": "",
                "text_blocks": [],
                "metadata": {},
                "error": f"解析失败: {str(e)}"
            }

    @classmethod
    def _parse_pdf(cls, file_path: Path) -> Dict:
        """解析PDF文件"""
        try:
            import pdfplumber

            content = []
            metadata = {}

            with pdfplumber.open(file_path) as pdf:
                # 提取元数据
                if pdf.metadata:
                    metadata = {
                        "title": pdf.metadata.get("Title", ""),
                        "author": pdf.metadata.get("Author", ""),
                        "pages": len(pdf.pages)
                    }

                # 提取文本
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        content.append(text)

            full_content = "\n\n".join(content)
            text_blocks = cls._split_into_blocks(full_content)

            return {
                "success": True,
                "content": full_content,
                "text_blocks": text_blocks,
                "metadata": metadata,
                "error": ""
            }

        except Exception as e:
            # 尝试使用PyPDF2作为备选
            return cls._parse_pdf_with_pypdf2(file_path)

    @classmethod
    def _parse_pdf_with_pypdf2(cls, file_path: Path) -> Dict:
        """使用PyPDF2解析PDF（备选方案）"""
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(str(file_path))
            content = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    content.append(text)

            full_content = "\n\n".join(content)
            text_blocks = cls._split_into_blocks(full_content)

            return {
                "success": True,
                "content": full_content,
                "text_blocks": text_blocks,
                "metadata": {"pages": len(reader.pages)},
                "error": ""
            }

        except Exception as e:
            return {
                "success": False,
                "content": "",
                "text_blocks": [],
                "metadata": {},
                "error": f"PDF解析失败: {str(e)}"
            }

    @classmethod
    def _parse_word(cls, file_path: Path) -> Dict:
        """解析Word文档"""
        try:
            from docx import Document

            doc = Document(str(file_path))
            content = []

            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)

            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content.append(" | ".join(row_text))

            full_content = "\n\n".join(content)
            text_blocks = cls._split_into_blocks(full_content)

            return {
                "success": True,
                "content": full_content,
                "text_blocks": text_blocks,
                "metadata": {"paragraphs": len(doc.paragraphs)},
                "error": ""
            }

        except Exception as e:
            return {
                "success": False,
                "content": "",
                "text_blocks": [],
                "metadata": {},
                "error": f"Word解析失败: {str(e)}"
            }

    @classmethod
    def _parse_text(cls, file_path: Path) -> Dict:
        """解析纯文本文件"""
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']
            content = None

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue

            if content is None:
                raise Exception("无法识别文件编码")

            text_blocks = cls._split_into_blocks(content)

            return {
                "success": True,
                "content": content,
                "text_blocks": text_blocks,
                "metadata": {"encoding": encoding},
                "error": ""
            }

        except Exception as e:
            return {
                "success": False,
                "content": "",
                "text_blocks": [],
                "metadata": {},
                "error": f"文本解析失败: {str(e)}"
            }

    @classmethod
    def _split_into_blocks(cls, text: str, max_length: int = 500) -> List[str]:
        """
        将文本分割成块，便于处理
        按段落分割，如果段落太长则按句子分割
        """
        blocks = []

        # 先按段落分割
        paragraphs = re.split(r'\n\s*\n', text)

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            if len(para) <= max_length:
                blocks.append(para)
            else:
                # 长段落按句子分割
                sentences = re.split(r'([。！？.!?]\s*)', para)
                current_block = ""

                for i in range(0, len(sentences), 2):
                    sentence = sentences[i]
                    if i + 1 < len(sentences):
                        sentence += sentences[i + 1]  # 加上标点

                    if len(current_block) + len(sentence) <= max_length:
                        current_block += sentence
                    else:
                        if current_block:
                            blocks.append(current_block.strip())
                        current_block = sentence

                if current_block:
                    blocks.append(current_block.strip())

        return blocks


class DocumentQAExtractor:
    """从文档中提取Q&A对"""

    @classmethod
    def extract_qa_pairs(cls, text_blocks: List[str]) -> List[Dict[str, str]]:
        """
        从文本块中提取问答对
        支持多种格式：
        1. Q: xxx A: xxx
        2. 问：xxx 答：xxx
        3. 问题：xxx 答案：xxx
        4. 数字编号的问题列表
        """
        qa_pairs = []

        # 合并文本以便正则匹配
        full_text = "\n".join(text_blocks)

        # 模式1: Q: xxx A: xxx
        pattern1 = r'Q[:：]\s*(.+?)\s*A[:：]\s*(.+?)(?=\s*Q[:：]|\s*$)'
        matches1 = re.findall(pattern1, full_text, re.DOTALL | re.IGNORECASE)
        for q, a in matches1:
            qa_pairs.append({
                "question": cls._clean_text(q),
                "answer": cls._clean_text(a)
            })

        # 模式2: 问：xxx 答：xxx
        pattern2 = r'问[:：]\s*(.+?)\s*答[:：]\s*(.+?)(?=\s*问[:：]|\s*$)'
        matches2 = re.findall(pattern2, full_text, re.DOTALL)
        for q, a in matches2:
            qa_pairs.append({
                "question": cls._clean_text(q),
                "answer": cls._clean_text(a)
            })

        # 模式3: 问题：xxx 答案：xxx
        pattern3 = r'问题[:：]\s*(.+?)\s*答案[:：]\s*(.+?)(?=\s*问题[:：]|\s*$)'
        matches3 = re.findall(pattern3, full_text, re.DOTALL)
        for q, a in matches3:
            qa_pairs.append({
                "question": cls._clean_text(q),
                "answer": cls._clean_text(a)
            })

        # 模式4: 数字编号的问题（如 1. xxx 或 1、xxx）
        pattern4 = r'(?:^|\n)\s*(?:\d+[.．、])\s*([^\n]+?)(?=\n\s*(?:\d+[.．、])|\n\s*(?:答|答案)[:：]|$)'
        questions = re.findall(pattern4, full_text)

        # 如果没有找到标准格式，尝试启发式提取
        if not qa_pairs and len(text_blocks) >= 2:
            # 假设奇数块是问题，偶数块是答案
            for i in range(0, len(text_blocks) - 1, 2):
                q = text_blocks[i]
                a = text_blocks[i + 1]

                # 判断是否是问答对（问题通常较短，以问号结尾）
                if cls._is_likely_question(q):
                    qa_pairs.append({
                        "question": cls._clean_text(q),
                        "answer": cls._clean_text(a)
                    })

        return qa_pairs

    @classmethod
    def _clean_text(cls, text: str) -> str:
        """清理文本"""
        # 移除多余空白
        text = re.sub(r'\s+', ' ', text)
        # 移除特殊字符
        text = text.strip()
        return text

    @classmethod
    def _is_likely_question(cls, text: str) -> bool:
        """判断文本是否可能是问题"""
        text = text.strip()

        # 以问号结尾
        if text.endswith('？') or text.endswith('?'):
            return True

        # 包含疑问词
        question_words = ['什么', '怎么', '为什么', '多少', '几', '是否', '能否', '可以吗']
        for word in question_words:
            if word in text:
                return True

        # 长度适中（问题通常不会太长）
        if 5 < len(text) < 100:
            return True

        return False
